#!/usr/bin/env python3
"""
Blog generator for date-of-death.com
Scans Google Drive Reports folder for Market Trends Report .docx files,
generates blog entry HTML pages, and builds paginated index pages.

Run daily at 4pm via scheduled task.
"""

import os
import sys
import glob
import json
import re
import csv
import math
import hashlib
from datetime import datetime, timedelta
from pathlib import Path

try:
    from docx import Document
except ImportError:
    os.system("pip install python-docx --break-system-packages -q")
    from docx import Document

# ── Configuration ──────────────────────────────────────────────────
REPORTS_DIR = os.environ.get(
    "REPORTS_DIR",
    "/Users/brianward/Library/CloudStorage/GoogleDrive-b@appraiser.llc/My Drive/Reports"
)
SITE_DIR = os.environ.get(
    "SITE_DIR",
    os.path.dirname(os.path.abspath(__file__))
)
BLOG_DIR = os.path.join(SITE_DIR, "blog")
PUBLISHED_FILE = os.path.join(BLOG_DIR, "published.json")
ENTRIES_PER_PAGE = 5
MAX_REPORT_AGE_DAYS = 7       # Only process reports generated in past N days
MAX_EFFECTIVE_AGE_DAYS = 60   # Only include if effective date within N days

# ── Helpers ────────────────────────────────────────────────────────

def slug_from_folder(folder_name):
    """Create a URL-friendly slug from a folder name."""
    s = folder_name.lower().strip()
    s = re.sub(r'[^a-z0-9]+', '-', s)
    s = s.strip('-')
    return s


def extract_location_from_csv(folder_path):
    """Extract city, zip codes, and market area from CSV files in the folder."""
    cities = {}  # city name (normalized) -> count
    zips = set()
    market_area = None

    # Prefer Property_System_Grid CSVs (the MLS data) over other CSV types
    csv_files = sorted(glob.glob(os.path.join(folder_path, "Property_System_Grid*.csv")))
    if not csv_files:
        csv_files = sorted(glob.glob(os.path.join(folder_path, "*.csv")))

    for csv_file in csv_files:
        try:
            with open(csv_file, 'r', errors='ignore') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    raw_city = row.get('City', row.get('city', '')).strip()
                    if raw_city:
                        # Normalize: title case, strip extra spaces
                        norm = raw_city.strip().title()
                        # Fix common MLS quirks
                        norm = re.sub(r'\s+', ' ', norm)
                        cities[norm] = cities.get(norm, 0) + 1
                    z = row.get('Zip', row.get('Zip Code', '')).strip()
                    if z:
                        # Keep just the 5-digit zip
                        z = z[:5] if len(z) >= 5 else z
                        zips.add(z)
                    if not market_area:
                        ma = row.get('Market Area', row.get('Area', row.get('Master Area', '')))
                        if ma and ma.strip():
                            market_area = ma.strip()
        except Exception:
            continue
        if cities:
            break  # Use first good CSV

    # Determine primary city by frequency (most common normalized name)
    primary_city = None
    if cities:
        # Group case variations and pick the most common
        primary_city = max(cities, key=cities.get)

    return primary_city, sorted(zips), market_area


def extract_address_from_folder(folder_name):
    """Parse a rough address from the folder name."""
    # Pattern: "XX NNNN Street Name" or "XX-NNNN-Street-Name-City-CA-XXXXX"
    name = folder_name.replace('-', ' ')
    # Remove leading abbreviation (LJ, SD, etc.)
    name = re.sub(r'^[A-Z]{2,3}\s+', '', name)
    return name


def _parse_date_str(date_str):
    """Try to parse a date string in common formats. Returns datetime or None."""
    for fmt in ('%m/%d/%Y', '%Y-%m-%d', '%B %d, %Y', '%B %d %Y'):
        try:
            return datetime.strptime(date_str.replace(',', '').strip(), fmt.replace(',', ''))
        except ValueError:
            continue
    return None


def extract_effective_date(docx_path, folder_path):
    """Extract the effective date from property.json, metrics.json, appraisal PDFs,
    invoice PDFs, or the .docx text in the folder.
    Returns a datetime object, or None if not found."""

    # Strategy 1: Check property.json (most reliable — written during appraisal setup)
    property_json_path = os.path.join(folder_path, 'property.json')
    if os.path.exists(property_json_path):
        try:
            with open(property_json_path, 'r') as f:
                pdata = json.load(f)
            eff = pdata.get('effectiveDate', '')
            if eff:
                dt = _parse_date_str(eff)
                if dt:
                    return dt
        except Exception:
            pass

    # Strategy 2: Check metrics.json in the same folder
    metrics_json_path = os.path.join(folder_path, 'metrics.json')
    if os.path.exists(metrics_json_path):
        try:
            with open(metrics_json_path, 'r') as f:
                mdata = json.load(f)
            eff = mdata.get('effective_date', '')
            if eff:
                dt = _parse_date_str(eff)
                if dt:
                    return dt
        except Exception:
            pass

    # Strategy 3: Check appraisal report PDFs (contain "Effective Date: MM/DD/YYYY")
    try:
        import fitz  # pymupdf
        # Check Appraisal*.pdf files first, then any other PDFs
        appraisal_pdfs = sorted(glob.glob(os.path.join(folder_path, "Appraisal*.pdf")))
        invoice_pdfs = [f for f in glob.glob(os.path.join(folder_path, "*.pdf"))
                        if 'invoice' in os.path.basename(f).lower()
                        or 'payment' in os.path.basename(f).lower()]
        for pdf_path in appraisal_pdfs + invoice_pdfs:
            try:
                pdf_doc = fitz.open(pdf_path)
                pdf_text = ''.join(page.get_text() for page in pdf_doc)
                m = re.search(r'Effective\s*Date[:\s]+(\d{1,2}/\d{1,2}/\d{4})', pdf_text, re.IGNORECASE)
                if m:
                    dt = _parse_date_str(m.group(1))
                    if dt:
                        return dt
            except Exception:
                continue
    except ImportError:
        pass  # pymupdf not available

    # Strategy 4: Parse from Market Trends .docx text
    try:
        doc = Document(docx_path)
        all_text = ' '.join(p.text for p in doc.paragraphs)
        patterns = [
            r'Effective\s+Date[:\s]+(\d{1,2}/\d{1,2}/\d{4})',
            r'effective\s+date\s+of\s+(\d{1,2}/\d{1,2}/\d{4})',
            r'ending\s+(\w+\s+\d{1,2},?\s+\d{4})',
            r'preceding\s+(\d{1,2}/\d{1,2}/\d{4})',
            r'preceding\s+(\w+\s+\d{1,2},?\s+\d{4})',
            r'as\s+of\s+(\d{1,2}/\d{1,2}/\d{4})',
            r'Effective\s+Date[:\s]+(\w+\s+\d{1,2},?\s+\d{4})',
        ]
        for pat in patterns:
            m = re.search(pat, all_text, re.IGNORECASE)
            if m:
                dt = _parse_date_str(m.group(1))
                if dt:
                    return dt
    except Exception:
        pass

    return None


def parse_report(docx_path):
    """Extract narrative text and key metrics from a Market Trends Report .docx."""
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Get narrative paragraphs (skip title and "Listing and Sales Data" / "Key Metrics" labels)
    skip_labels = [
        'market trends report',
        'listing and sales data',
        'key metrics at a glance',
        'summary and conclusions',
    ]
    narrative = []
    summary = []
    in_summary = False
    for p in paragraphs:
        if p.lower().strip() in skip_labels:
            if 'summary' in p.lower():
                in_summary = True
            continue
        if in_summary:
            summary.append(p)
        elif not any(p.lower().startswith(s) for s in skip_labels):
            if not in_summary and p != paragraphs[0]:
                narrative.append(p)

    # Extract key metrics from table
    metrics = {}
    if len(doc.tables) >= 2:
        table = doc.tables[1]
        for row in table.rows[1:]:  # Skip header
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 3:
                metrics[cells[0]] = {'value': cells[1], 'trend': cells[2]}

    return {
        'narrative': narrative,
        'summary': summary,
        'metrics': metrics,
    }


def generate_blog_entry_html(entry):
    """Generate a full HTML page for a single blog entry."""
    title = entry['title']
    date_str = entry['date_display']
    city = entry['city']
    market_area = entry['market_area']
    data = entry['data']
    slug = entry['slug']

    # Build metrics table rows
    metrics_rows = ""
    for key, val in data['metrics'].items():
        trend_class = ""
        if val['trend'].lower() in ['increasing', 'positive', 'strengthening', 'improving']:
            trend_class = "trend-up"
        elif val['trend'].lower() in ['declining', 'negative', 'decreasing', 'worsening']:
            trend_class = "trend-down"
        else:
            trend_class = "trend-stable"
        metrics_rows += f'<tr><td>{key}</td><td><strong>{val["value"]}</strong></td><td class="{trend_class}">{val["trend"]}</td></tr>\n'

    # Build summary paragraphs
    summary_html = ""
    for p in data['summary']:
        summary_html += f"<p>{p}</p>\n"
    if not summary_html:
        # Use first narrative paragraph as summary
        for p in data['narrative'][:2]:
            summary_html += f"<p>{p}</p>\n"

    zip_codes = entry.get('zip_codes', [])
    zip_label = ", ".join(zip_codes) if zip_codes else ""
    if zip_label:
        location_label = f"{city} {zip_label}"
    else:
        location_label = city

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} | Date-of-Death Appraisals</title>
    <meta name="description" content="Real estate market trends analysis for {location_label}, California. Median prices, inventory levels, and market conditions for estate and date-of-death valuations.">
    <meta name="robots" content="index, follow">
    <meta property="og:title" content="{title}">
    <meta property="og:description" content="Real estate market trends for {location_label} &mdash; key metrics for estate valuations.">
    <meta property="og:type" content="article">
    <meta property="og:url" content="https://date-of-death.com/blog/{slug}">
    <meta name="twitter:card" content="summary">
    <meta name="twitter:title" content="{title}">
    <meta name="twitter:description" content="Real estate market trends for {location_label} &mdash; key metrics for estate valuations.">
    <link rel="canonical" href="https://date-of-death.com/blog/{slug}">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="/styles.css?v=5">
    <script type="application/ld+json">
    {{
        "@context": "https://schema.org",
        "@type": "Article",
        "headline": "{title}",
        "datePublished": "{entry['date_iso']}",
        "author": {{
            "@type": "Person",
            "name": "Brian Ward",
            "jobTitle": "California Certified Residential Appraiser"
        }},
        "publisher": {{
            "@type": "Organization",
            "name": "Date-of-Death Appraisals",
            "url": "https://date-of-death.com"
        }}
    }}
    </script>
</head>
<body>

<header class="site-header">
    <div class="header-inner">
        <a href="/" class="logo">Date-of-Death Appraisals</a>
        <button class="menu-toggle" aria-label="Toggle navigation" onclick="document.querySelector('.main-nav').classList.toggle('open')">
            <span></span><span></span><span></span>
        </button>
        <nav class="main-nav">
            <a href="/why-desktop">Why Desktop</a>
            <a href="/how-it-works">How It Works</a>
            <a href="/pricing">Pricing</a>
            <a href="/faq">FAQ</a>
            <a href="/about">About</a>
            <a href="/contact">Contact</a>
            <a href="/service-areas">Service Areas</a>
            <a href="/order" class="nav-cta">Order Now</a>
        </nav>
    </div>
</header>

<section class="hero hero-blog">
    <div class="container">
        <h1>Market Trends</h1>
        <p class="subtitle">Real estate market analysis from our appraisal practice across California.</p>
    </div>
</section>

<section>
    <div class="container blog-content">
        <p class="blog-meta"><a href="/blog">&larr; All Posts</a> &nbsp;&middot;&nbsp; {date_str} &nbsp;&middot;&nbsp; {location_label}, California</p>
        <h2>{title}</h2>

        <div class="blog-metrics-table">
            <h3>Key Metrics at a Glance</h3>
            <table>
                <thead>
                    <tr><th>Metric</th><th>Value</th><th>Trend</th></tr>
                </thead>
                <tbody>
                    {metrics_rows}
                </tbody>
            </table>
        </div>

        <h3>Analysis</h3>
        {summary_html}

        <div class="blog-cta-box">
            <h3>Need a Date-of-Death Appraisal?</h3>
            <p>Our market trend analyses are performed as part of every appraisal we prepare. If you need a USPAP-compliant date-of-death appraisal for the IRS step-up in basis, we provide certified desktop appraisals statewide across California &mdash; starting at $299.</p>
            <a href="/order" class="hero-cta">Order Your Appraisal</a>
        </div>
    </div>
</section>

<footer class="site-footer">
    <div class="footer-inner">
        <div class="footer-col">
            <h4>Services</h4>
            <a href="/pricing">Pricing</a>
            <a href="/why-desktop">Why Desktop Appraisals</a>
            <a href="/how-it-works">How It Works</a>
            <a href="/order">Order Now</a>
        </div>
        <div class="footer-col">
            <h4>Resources</h4>
            <a href="/faq">FAQ</a>
            <a href="/about">About &amp; Credentials</a>
            <a href="/service-areas">Service Areas</a>
            <a href="/step-up-in-basis">Step-Up in Basis Guide</a>
            <a href="/blog">Market Trends</a>
            <a href="/reviews">Reviews</a>
        </div>
        <div class="footer-col">
            <h4>Contact</h4>
            <a href="/contact">Contact Form</a>
        </div>
    </div>
    <div class="footer-bottom">
        <span>&copy; 2026 Brian Ward Appraisal. All rights reserved.</span>
        <span>California Certified Residential Appraiser &middot; License No. AR036053</span>
    </div>
</footer>

</body>
</html>"""


def generate_index_page(entries, page_num, total_pages):
    """Generate a paginated blog index page."""
    is_first = page_num == 1
    path_prefix = "/blog" if is_first else f"/blog/page/{page_num}"

    cards_html = ""
    for entry in entries:
        zip_codes = entry.get('zip_codes', [])
        zip_label = ", ".join(zip_codes) if zip_codes else ""
        if zip_label:
            location = f"{entry['city']} {zip_label}"
        else:
            location = entry['city']

        # Short summary for the card
        summary_text = ""
        if entry['data']['summary']:
            summary_text = entry['data']['summary'][0][:200]
        elif entry['data']['narrative']:
            summary_text = entry['data']['narrative'][0][:200]
        if len(summary_text) >= 200:
            summary_text = summary_text[:197] + "..."

        # Key metric highlights
        metrics = entry['data']['metrics']
        price_val = metrics.get('Median Sale Price (0-3 Mo)', {}).get('value', '')
        ppsf_val = metrics.get('Median $/SF (0-3 Mo)', {}).get('value', '')
        inventory_val = metrics.get('Months of Inventory', {}).get('value', '')

        cards_html += f"""
        <article class="blog-card">
            <div class="blog-card-header">
                <span class="blog-card-date">{entry['date_display']}</span>
                <span class="blog-card-location">{location}, CA</span>
            </div>
            <h3><a href="/blog/{entry['slug']}">{entry['title']}</a></h3>
            <div class="blog-card-metrics">
                <span>Median Price: <strong>{price_val}</strong></span>
                <span>$/SF: <strong>{ppsf_val}</strong></span>
                <span>Inventory: <strong>{inventory_val} mo</strong></span>
            </div>
            <p>{summary_text}</p>
            <a href="/blog/{entry['slug']}" class="blog-read-more">Read full analysis &rarr;</a>
        </article>"""

    # Pagination
    pagination_html = ""
    if total_pages > 1:
        pagination_html = '<div class="blog-pagination">'
        if page_num > 1:
            prev_url = "/blog" if page_num == 2 else f"/blog/page/{page_num - 1}"
            pagination_html += f'<a href="{prev_url}" class="blog-page-link">&larr; Newer</a>'
        for p in range(1, total_pages + 1):
            url = "/blog" if p == 1 else f"/blog/page/{p}"
            active = ' class="blog-page-active"' if p == page_num else ''
            pagination_html += f'<a href="{url}"{active}>{p}</a>'
        if page_num < total_pages:
            pagination_html += f'<a href="/blog/page/{page_num + 1}" class="blog-page-link">Older &rarr;</a>'
        pagination_html += '</div>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Market Trends Blog{'' if is_first else f' — Page {page_num}'} | Date-of-Death Appraisals</title>
    <meta name="description" content="California real estate market trends and analysis from our appraisal practice. Median prices, inventory, and market conditions across California communities.">
    <meta name="robots" content="index, follow">
    <meta property="og:title" content="Market Trends Blog | Date-of-Death Appraisals">
    <meta property="og:description" content="Real estate market analysis from our appraisal practice across California.">
    <meta property="og:type" content="website">
    <meta property="og:url" content="https://date-of-death.com{path_prefix}">
    <meta name="twitter:card" content="summary">
    <meta name="twitter:title" content="Market Trends Blog | Date-of-Death Appraisals">
    <meta name="twitter:description" content="Real estate market analysis from our appraisal practice across California.">
    <link rel="canonical" href="https://date-of-death.com{path_prefix}">
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="/styles.css?v=5">
</head>
<body>

<header class="site-header">
    <div class="header-inner">
        <a href="/" class="logo">Date-of-Death Appraisals</a>
        <button class="menu-toggle" aria-label="Toggle navigation" onclick="document.querySelector('.main-nav').classList.toggle('open')">
            <span></span><span></span><span></span>
        </button>
        <nav class="main-nav">
            <a href="/why-desktop">Why Desktop</a>
            <a href="/how-it-works">How It Works</a>
            <a href="/pricing">Pricing</a>
            <a href="/faq">FAQ</a>
            <a href="/about">About</a>
            <a href="/contact">Contact</a>
            <a href="/service-areas">Service Areas</a>
            <a href="/order" class="nav-cta">Order Now</a>
        </nav>
    </div>
</header>

<section class="hero hero-blog">
    <div class="container">
        <h1>Market Trends</h1>
        <p class="subtitle">Real estate market analysis from our appraisal practice across California.</p>
    </div>
</section>

<section>
    <div class="container blog-list">
        {cards_html}
        {pagination_html}
    </div>
</section>

<section class="cta-banner">
    <h2>Need a Date-of-Death Appraisal?</h2>
    <p>USPAP-compliant desktop appraisals for the IRS step-up in basis &mdash; serving all 58 California counties.</p>
    <a href="/order" class="hero-cta">Order Your Appraisal &mdash; From $299</a>
</section>

<footer class="site-footer">
    <div class="footer-inner">
        <div class="footer-col">
            <h4>Services</h4>
            <a href="/pricing">Pricing</a>
            <a href="/why-desktop">Why Desktop Appraisals</a>
            <a href="/how-it-works">How It Works</a>
            <a href="/order">Order Now</a>
        </div>
        <div class="footer-col">
            <h4>Resources</h4>
            <a href="/faq">FAQ</a>
            <a href="/about">About &amp; Credentials</a>
            <a href="/service-areas">Service Areas</a>
            <a href="/step-up-in-basis">Step-Up in Basis Guide</a>
            <a href="/blog">Market Trends</a>
            <a href="/reviews">Reviews</a>
        </div>
        <div class="footer-col">
            <h4>Contact</h4>
            <a href="/contact">Contact Form</a>
        </div>
    </div>
    <div class="footer-bottom">
        <span>&copy; 2026 Brian Ward Appraisal. All rights reserved.</span>
        <span>California Certified Residential Appraiser &middot; License No. AR036053</span>
    </div>
</footer>

</body>
</html>"""


# ── Main ───────────────────────────────────────────────────────────

def main():
    now = datetime.now()
    cutoff_generated = now - timedelta(days=MAX_REPORT_AGE_DAYS)
    cutoff_effective = now - timedelta(days=MAX_EFFECTIVE_AGE_DAYS)

    # Load published entries tracker
    os.makedirs(BLOG_DIR, exist_ok=True)
    published = {}
    if os.path.exists(PUBLISHED_FILE):
        with open(PUBLISHED_FILE, 'r') as f:
            published = json.load(f)

    # Find all Market Trends Report .docx files
    patterns = [
        os.path.join(REPORTS_DIR, "**", "Market*Trends*Report*.docx"),
        os.path.join(REPORTS_DIR, "**", "Market_Trends_Report*.docx"),
    ]
    found_files = set()
    for pat in patterns:
        for f in glob.glob(pat, recursive=True):
            found_files.add(f)

    # Filter and process reports
    new_entries = []
    for filepath in sorted(found_files):
        folder_path = os.path.dirname(filepath)
        folder_name = os.path.basename(folder_path)

        # Skip root-level test files and "files" folders
        if folder_name in ('Reports', 'files', ''):
            continue
        # Skip files with "(1)" duplicates
        if '(1)' in os.path.basename(filepath):
            continue

        mod_time = datetime.fromtimestamp(os.path.getmtime(filepath))

        # Filter: generated within past week (based on file modification time)
        if mod_time < cutoff_generated:
            continue

        # Extract the actual effective date from the report
        effective_date = extract_effective_date(filepath, folder_path)
        if effective_date is None:
            # Fallback: use file modification date only if effective date not found
            effective_date = mod_time
            print(f"  WARNING: No effective date found in {folder_name}, using mtime")

        # Filter: effective date within 60 days (use actual effective date, not mtime)
        if effective_date < cutoff_effective:
            continue

        # Duplicate check: use folder_name as unique key
        # If already published with same or newer modification time, skip
        entry_key = slug_from_folder(folder_name)
        if entry_key in published:
            prev_mod = published[entry_key].get('mod_time', '')
            if prev_mod >= mod_time.isoformat():
                print(f"  SKIP (already published): {folder_name}")
                continue

        # Extract city, zip codes, and market area from CSV
        city, zip_codes, market_area = extract_location_from_csv(folder_path)
        if not city:
            # Try to parse from folder name
            addr = extract_address_from_folder(folder_name)
            # Common abbreviations
            abbrevs = {
                'LJ': 'La Jolla', 'SD': 'San Diego', 'OC': 'Oceanside',
                'EN': 'Encinitas', 'FB': 'Fallbrook', 'LM': 'La Mesa',
                'RM': 'Rancho Mirage', 'CV': 'Chula Vista', 'RB': 'Rancho Bernardo',
                'SC': 'Santa Clarita', 'SM': 'San Marcos', 'PS': 'Palm Springs', 'DA': 'Davis',
            }
            prefix = folder_name[:2]
            city = abbrevs.get(prefix, addr.split()[0] if addr else 'California')
        if not zip_codes:
            # Try to extract zip from folder name (e.g. "SD-18520-...-CA-92128-1109")
            folder_zips = re.findall(r'\b9[0-9]{4}\b', folder_name)
            zip_codes = sorted(set(folder_zips))

        # Parse the report
        try:
            data = parse_report(filepath)
        except Exception as e:
            print(f"  ERROR parsing {filepath}: {e}")
            continue

        if not data['metrics']:
            print(f"  SKIP (no metrics): {folder_name}")
            continue

        # Build title with city and zip codes (use effective date for the month/year)
        zip_label = ", ".join(zip_codes) if zip_codes else ""
        if zip_label:
            title = f"Market Trends: {city} {zip_label} — {effective_date.strftime('%B %Y')}"
        else:
            title = f"Market Trends: {city} — {effective_date.strftime('%B %Y')}"

        entry = {
            'slug': entry_key,
            'title': title,
            'city': city,
            'zip_codes': zip_codes,
            'market_area': market_area or '',
            'date_iso': effective_date.strftime('%Y-%m-%d'),
            'date_display': effective_date.strftime('%B %d, %Y'),
            'mod_time': mod_time.isoformat(),
            'folder_name': folder_name,
            'data': data,
        }
        new_entries.append(entry)
        print(f"  NEW: {title}")

    # Merge with existing published entries (keep old entries that are still valid)
    all_entries = {}

    # Load existing entries data
    for key, meta in published.items():
        all_entries[key] = meta

    # Add/update new entries
    for entry in new_entries:
        # Write individual blog entry page
        entry_dir = os.path.join(BLOG_DIR, entry['slug'])
        os.makedirs(entry_dir, exist_ok=True)
        html = generate_blog_entry_html(entry)
        with open(os.path.join(entry_dir, 'index.html'), 'w') as f:
            f.write(html)

        # Store metadata (without full data for the JSON tracker)
        all_entries[entry['slug']] = {
            'slug': entry['slug'],
            'title': entry['title'],
            'city': entry['city'],
            'zip_codes': entry.get('zip_codes', []),
            'market_area': entry['market_area'],
            'date_iso': entry['date_iso'],
            'date_display': entry['date_display'],
            'mod_time': entry['mod_time'],
            'folder_name': entry['folder_name'],
            'metrics_summary': {k: v for k, v in entry['data']['metrics'].items()},
            'summary_text': entry['data']['summary'][0][:300] if entry['data']['summary'] else (entry['data']['narrative'][0][:300] if entry['data']['narrative'] else ''),
        }

    # Save published tracker
    with open(PUBLISHED_FILE, 'w') as f:
        json.dump(all_entries, f, indent=2)

    # ── Rebuild all index pages from published data ──
    # Sort all entries by date (newest first)
    sorted_entries = sorted(
        all_entries.values(),
        key=lambda e: e.get('date_iso', ''),
        reverse=True
    )

    # For index pages, we need the full data. Re-read entries that exist on disk.
    index_entries = []
    for meta in sorted_entries:
        slug = meta['slug']
        entry_html_path = os.path.join(BLOG_DIR, slug, 'index.html')
        if os.path.exists(entry_html_path):
            # Reconstruct minimal entry for index card
            metrics = meta.get('metrics_summary', {})
            summary = meta.get('summary_text', '')
            index_entries.append({
                'slug': slug,
                'title': meta['title'],
                'city': meta['city'],
                'zip_codes': meta.get('zip_codes', []),
                'market_area': meta.get('market_area', ''),
                'date_iso': meta['date_iso'],
                'date_display': meta['date_display'],
                'data': {
                    'metrics': metrics,
                    'summary': [summary] if summary else [],
                    'narrative': [],
                },
            })

    total_pages = max(1, math.ceil(len(index_entries) / ENTRIES_PER_PAGE))

    # Generate page 1 (blog/index.html)
    for page_num in range(1, total_pages + 1):
        start = (page_num - 1) * ENTRIES_PER_PAGE
        end = start + ENTRIES_PER_PAGE
        page_entries = index_entries[start:end]

        html = generate_index_page(page_entries, page_num, total_pages)

        if page_num == 1:
            with open(os.path.join(BLOG_DIR, 'index.html'), 'w') as f:
                f.write(html)
        else:
            page_dir = os.path.join(BLOG_DIR, 'page', str(page_num))
            os.makedirs(page_dir, exist_ok=True)
            with open(os.path.join(page_dir, 'index.html'), 'w') as f:
                f.write(html)

    print(f"\nBlog updated: {len(index_entries)} total entries, {total_pages} pages")
    if new_entries:
        print(f"  New entries added: {len(new_entries)}")
    else:
        print("  No new entries to add")


if __name__ == '__main__':
    main()
